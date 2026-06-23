import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Титульный лист
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЕТ ПО ИНДИВИДУАЛЬНОМУ ЗАДАНИЮ №11\n")
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph("Дисциплина: Информационная безопасность\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.add_run(f"\n\nСтудент: Артыков Амирбек\nГруппа: БИС 3-24\nНаправление: Информационная безопасность")

    doc.add_page_break()

    # 1. Описание
    doc.add_heading('1. Общие сведения о проекте', level=1)
    doc.add_paragraph(
        "Веб-приложение предназначено для инвентаризации сетевого оборудования. "
        "Разработка велась с соблюдением методологии Secure SDLC. Основные функции включают "
        "учет устройств, управление интерфейсами и детальный аудит действий администраторов."
    )

    # 2. Стек
    doc.add_heading('2. Технологический стек', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Слой'
    hdr_cells[1].text = 'Технологии'
    
    tech_data = [
        ('Backend', 'ASP.NET Core 9, EF Core, SQLite'),
        ('Frontend', 'React 18, Vite 6, Axios'),
        ('Безопасность', 'JWT (HttpOnly), BCrypt, RBAC, CORS')
    ]
    for layer, tech in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech

    # 3. Безопасность
    doc.add_heading('3. Реализованные механизмы защиты', level=1)
    security_points = [
        ("JWT в HttpOnly Cookie", "Защита от кражи токена через XSS-атаки. JavaScript не имеет доступа к cookie."),
        ("RBAC", "Разграничение прав: Admin (полный доступ), User (только чтение)."),
        ("Аудит (Accountability)", "Автоматическая фиксация автора и времени изменений в AppDbContext."),
        ("Защита от SQLi", "Параметризация запросов через Entity Framework Core."),
        ("DTO", "Использование объектов передачи данных для предотвращения атак типа Mass Assignment (Over-posting).")
    ]
    for title, desc in security_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # 4. Архитектура
    doc.add_heading('4. Архитектура системы', level=1)
    doc.add_paragraph(
        "Приложение построено по классической клиент-серверной архитектуре. "
        "Frontend взаимодействует с Backend через REST API. "
        "Безопасность обеспечивается Middleware-компонентами на стороне сервера, "
        "проверяющими валидность JWT при каждом запросе."
    )

    # 5. Код аудита
    doc.add_heading('5. Пример кода: Механизм аудита', level=1)
    code = (
        "// Фрагмент из AppDbContext.cs\n"
        "foreach (var entry in ChangeTracker.Entries<Device>())\n"
        "{\n"
        "    if (entry.State is EntityState.Added or EntityState.Modified)\n"
        "    {\n"
        "        entry.Entity.LastModifiedDate = DateTime.UtcNow;\n"
        "        entry.Entity.LastModifiedByUserId = currentUserId;\n"
        "    }\n"
        "}"
    )
    doc.add_paragraph(code, style='Code') # Предполагается наличие стиля или просто моноширинный текст

    file_path = "Report_Artykov_Amirbek.docx"
    doc.save(file_path)
    print(f"Отчет успешно создан: {file_path}")

if __name__ == "__main__":
    create_report()